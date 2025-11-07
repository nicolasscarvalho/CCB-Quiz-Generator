import streamlit as st
import os
import random
import json
from docx import Document
from docx.shared import Pt
import io
import time
import re

# --- I18N (Internationalization) Setup ---

# Dicionário de traduções para Português (Brasil)
pt_BR = {
    # Page Config
    "page_title": "Gerador de Provas de Inglês",
    "page_icon": "📄",
    
    # Setup/Error Messages
    "warn_books_dir_not_found": f"O diretório `BOOKS` não foi encontrado.",
    "btn_create_sample_structure": "Clique aqui para criar uma estrutura de pastas e arquivos .json de amostra",
    "info_creating_env": "Criando ambiente de teste com arquivos .json... Por favor, recarregue a página em alguns segundos.",
    
    # Sidebar
    "sidebar_lang_title": "Idioma",
    "sidebar_lang_label": "Escolha o idioma:",
    "sidebar_header": "📖 Configuração da Prova",
    "sb_select_book": "Escolha o Livro",
    "err_no_books": "Nenhum livro encontrado no diretório '{}'.",
    "err_no_units": "Nenhuma unidade encontrada para o livro '{}'. Por favor, verifique a estrutura das pastas.",
    "sb_select_units_title": "🎯 Selecione as Unidades Desejadas",
    "sb_select_units_label": "Selecione as Unidades (1, 2, 3, 4, 5, 6)",
    "sb_q_config_title": "⚙️ Número de Questões por Seção",
    "sb_q_config_grammar": "Gramática",
    "sb_q_config_vocab": "Vocabulário",
    "sb_total_questions": "**Total de questões na prova: {total}**",
    
    # Main Page
    "warn_no_unit_selected": "Por favor, selecione pelo menos uma unidade na barra lateral para continuar.",
    "main_summary_title": "Resumo da Configuração",
    "main_summary_book": "**📖 Livro**",
    "main_summary_units": "**📚 Unidades Selecionadas**",
    "main_summary_q_config": "**⚙️ Questões por Seção**",
    "main_q_summary_content": "G: {grammar} | V: {vocabulary}",
    "btn_generate_std": "🚀 Gerar Prova Padrão",
    "btn_generate_custom": "🚀 Gerar Prova Personalizada",
    "spinner_generating": "Lendo arquivos JSON e montando sua prova...",
    "err_generation": "Ocorreu um erro durante a geração da prova: {error}",
    "warn_no_questions_selected": "Por favor, selecione pelo menos uma questão para gerar a prova.",
    "btn_download": "📥 Baixar Prova Gerada (DOCX)",
    "filename_test": "Prova_{book}_Unidades_{units}.docx",
    
    # Docx Generation
    "docx_title": "Prova de Inglês - Livro: {book} | Unidade(s): {units}",
    "docx_name_date": "Nome: __________________________________________________ Data: ___/___/______",
    "docx_no_questions_found": "Nenhuma questão foi encontrada com os critérios selecionados.",
    "docx_section_header": "Seção: {section}",
    "docx_answer_key_title": "Gabarito (Uso do Professor)",
    "docx_answer_key_question": "Questão {number}:",

    # About Section
    "about_title": "Sobre o Gerador de Provas do CCB",
    "about_p1": "Professores da Casa de Cultura Britânica passavam horas criando provas manualmente. Criamos um sistema que gera essas mesmas experiências em segundos, liberando-os para focar no que realmente importa: ensinar. Além disso, imagine ter acesso a exercícios de inglês constantemente atualizados com base exatamente no que você está estudando? Esse processo era manual e lento, então o automatizamos para que alunos e professores tenham acesso a materiais de qualidade com o clique de um botão.",
    "about_p2": "Nosso projeto é uma aplicação web que gera questões de inglês personalizadas sobre conteúdo de gramática e vocabulário com base nos livros da coleção English File. Nossa tecnologia atua como um professor especialista, usando IA para ler materiais de aprendizado de nosso banco de dados e criar milhares de variações de questões baseadas unicamente em nossas fontes confiáveis.",
    "contribute_title": "Por favor, avalie-nos e contribua",
    "contribute_link_text": "[Clique aqui]({link}) para que possamos continuar melhorando esta ferramenta e garantir que ela permaneça precisa, rápida e simples. Sua experiência de usuário é essencial. Seu feedback nos permite:",
    "contribute_li1": "**Validar a Qualidade**: Garantir que o conteúdo gerado esteja alinhado às suas necessidades e à estrutura dos livros;",
    "contribute_li2": "**Priorizar Melhorias**: Entender onde investir nosso tempo de desenvolvimento, seja na otimização da geração de questões ou na usabilidade da interface;",
    "contribute_li3": "**Manter o Acesso Gratuito**: Sua participação valida a importância deste projeto para a comunidade da UFC.",
    "footer_text": "Desenvolvido com ❤️ por alunos da UFC"
}

# Dicionário de traduções para Inglês Britânico
en_GB = {
    # Page Config
    "page_title": "English Test Generator",
    "page_icon": "📄",

    # Setup/Error Messages
    "warn_books_dir_not_found": f"The `BOOKS` directory was not found.",
    "btn_create_sample_structure": "Click here to create a sample folder structure and .json files",
    "info_creating_env": "Creating test environment with .json files... Please reload the page in a few seconds.",

    # Sidebar
    "sidebar_lang_title": "Language",
    "sidebar_lang_label": "Choose language:",
    "sidebar_header": "📖 Test Configuration",
    "sb_select_book": "Choose Book",
    "err_no_books": "No books found in the '{}' directory.",
    "err_no_units": "No units found for the book '{}'. Please check the folder structure.",
    "sb_select_units_title": "🎯 Select Desired Units",
    "sb_select_units_label": "Select Units (1, 2, 3, 4, 5, 6)",
    "sb_q_config_title": "⚙️ Number of Questions per Section",
    "sb_q_config_grammar": "Grammar",
    "sb_q_config_vocab": "Vocabulary",
    "sb_total_questions": "**Total questions on the test: {total}**",

    # Main Page
    "warn_no_unit_selected": "Please select at least one unit from the sidebar to continue.",
    "main_summary_title": "Configuration Summary",
    "main_summary_book": "**📖 Book**",
    "main_summary_units": "**📚 Selected Units**",
    "main_summary_q_config": "**⚙️ Questions per Section**",
    "main_q_summary_content": "G: {grammar} | V: {vocabulary}",
    "btn_generate_std": "🚀 Generate Standard Test",
    "btn_generate_custom": "🚀 Generate Custom Test",
    "spinner_generating": "Reading JSON files and assembling your test...",
    "err_generation": "An error occurred during test generation: {error}",
    "warn_no_questions_selected": "Please select at least one question to generate the test.",
    "btn_download": "📥 Download Generated Test (DOCX)",
    "filename_test": "Test_{book}_Units_{units}.docx",

    # Docx Generation
    "docx_title": "English Test - Book: {book} | Unit(s): {units}",
    "docx_name_date": "Name: __________________________________________________ Date: ___/___/______",
    "docx_no_questions_found": "No questions were found with the selected criteria.",
    "docx_section_header": "Section: {section}",
    "docx_answer_key_title": "Answer Key (For Teacher's Use)",
    "docx_answer_key_question": "Question {number}:",

    # About Section
    "about_title": "About CCB's Quiz Generator",
    "about_p1": "Teachers at the Casa de Cultura Britânica spent hours manually creating tests. We created a system that generates these same experiences in seconds, freeing them to focus on what really matters: teaching. Furthermore, imagine having access to constantly updated English exercises based exactly on what you're studying? This process used to be manual and slow, so we've automated it so students and teachers have access to quality materials at the click of a button.",
    "about_p2": "Our project is a web application that generates customised English questions about grammar and vocabulary content based on books in the English File collection. Our technology acts like an expert teacher, using AI to read learning materials from our database and create thousands of question variations based solely on our trusted sources.",
    "contribute_title": "Please rate us and contribute",
    "contribute_link_text": "[Click here]({link}) so we can continue improving this tool and ensure it remains accurate, fast, and simple. Your user experience is essential. Your feedback allows us to:",
    "contribute_li1": "**Validate Quality**: Ensure that the generated content aligns with your needs and the structure of the books;",
    "contribute_li2": "**Prioritise Improvements**: Understand where to invest our development time, whether in optimising question generation or interface usability;",
    "contribute_li3": "**Maintain Free Access**: Your participation validates the importance of this project for the UFC community.",
    "footer_text": "Developed with ❤️ by UFC students"
}

LANGUAGES = {"pt_BR": pt_BR, "en_GB": en_GB}
LANG_OPTIONS_DISPLAY = {"Português (Brasil)": "pt_BR", "English (UK)": "en_GB"}

# Define o idioma padrão se ainda não estiver definido
if "lang" not in st.session_state:
    st.session_state.lang = "pt_BR" # Você pode mudar o padrão aqui para 'en_GB'

def get_lang(key: str) -> str:
    """Busca uma string de tradução com base no idioma atual no session_state."""
    lang_code = st.session_state.get("lang", "pt_BR")
    # Usa en_GB como fallback se a chave não for encontrada no idioma selecionado
    return LANGUAGES.get(lang_code, en_GB).get(key, LANGUAGES["en_GB"].get(key, f"Missing_Key: {key}"))

# --- Fim do Setup I18N ---


# --- Page Setup (AGORA USA A FUNÇÃO get_lang) ---
st.set_page_config(
    page_title=get_lang("page_title"),
    page_icon=get_lang("page_icon"),
    layout="wide"
)

# --- Logic Constants and Functions ---
BOOKS_DIR = "BOOKS"
SECTIONS = ["GRAMMAR", "VOCABULARY"]
DEFAULT_QUESTIONS = {"grammar": 3, "vocabulary": 3}
EVALUATION_LINK = "https://docs.google.com/forms/d/e/1FAIpQLSdm1n218RAyl_js-lQGvbWd_voBJlu_wZ90T_9p5dBaatD6Ew/viewform?usp=header"

@st.cache_data
def get_available_books(directory: str) -> list:
    """Returns a list of directories (books) inside the main directory."""
    if not os.path.exists(directory): return []
    return [d for d in os.listdir(directory) if os.path.isdir(os.path.join(directory, d))]

@st.cache_data
def parse_available_units(book_name: str) -> dict:
    """
    Analyzes the folder structure and extracts all unique unit designations 
    (1A, 1B, 2C) from the JSON 'topic' field.
    """
    if not book_name: return {}
    parsed_units = {}
    book_path = os.path.join(BOOKS_DIR, book_name)
    if not os.path.exists(book_path):
        return {}
    
    for unit_folder_name in os.listdir(book_path):
        if unit_folder_name.startswith("UNIT-"):
            unit_path = os.path.join(book_path, unit_folder_name)
            if os.path.isdir(unit_path):
                for filename in os.listdir(unit_path):
                    if filename.endswith(".json"):
                        file_path = os.path.join(unit_path, filename)
                        try:
                            with open(file_path, 'r', encoding='utf-8') as f:
                                data = json.load(f)
                                for q in data.get("questions", []):
                                    for topic_unit in q.get("topic", []):
                                        match = re.match(r"(\d+)([A-Za-z]*)", topic_unit)
                                        if match:
                                            num_part, alpha_part = match.groups()
                                            if num_part not in parsed_units:
                                                parsed_units[num_part] = set()
                                            if alpha_part:
                                                parsed_units[num_part].add(alpha_part.upper())
                                            elif not alpha_part:
                                                parsed_units[num_part] = parsed_units.get(num_part, set())
                        except Exception:
                            continue
    
    for num_part in parsed_units:
        parsed_units[num_part] = sorted(list(parsed_units[num_part]))
        
    return parsed_units


def write_question_to_doc(doc, question_data, question_number):
    """
    Formats and writes a single question to the docx document based on its type.
    """
    # --- 1. Write Header and Instructions ---
    q_instructions = question_data.get("instructions", "")
    q_type = question_data.get("type")

    # Write question number (main question) and instructions
    p_question_num = doc.add_paragraph()
    p_question_num.add_run(f"{question_number}. ").bold = True
    p_question_num.add_run(q_instructions).bold = True
    doc.add_paragraph() 

    # Write example if it exists
    example_data = question_data.get("example")
    if isinstance(example_data, dict):
        example_item = example_data.get("item", "")
        doc.add_paragraph(f"Example: {example_item}")
    elif isinstance(example_data, str):
        doc.add_paragraph(f"Example: {example_data}")

    qa_pairs = question_data.get("qa_pairs") or question_data.get("qa_pair", [])

    # --- 2. Conditional Formatting based on Question Type ---
    
    # Type A: Tick/Select Sentence (A/B comparison, elimina quebra de linha)
    if q_type in ["tick_correct_sentence", "select_correct_sentence"]:
        for i, pair in enumerate(qa_pairs):
            item_text = pair.get("item", "") 
            options = item_text.split(' / ')
            options_text_list = []
            
            for j, option in enumerate(options):
                if j == 0 and re.match(r"^[A-D]\s", option):
                    options_text_list.append(f"( ) {option}")
                else:
                    options_text_list.append(option)
            
            consolidated_text = f"{i+1}. {' '.join(options_text_list)}"
            doc.add_paragraph(consolidated_text)
            doc.add_paragraph() # Espaçamento entre sub-itens
            
    # Type B: Underline/Select Word
    elif q_type.startswith(("underline_correct_word", "select_correct_possessive_adjective", 
                            "underline_correct_word_subject_possessive", "underline_correct_word_or_phrase")):
        for i, pair in enumerate(qa_pairs):
            item_text = pair.get("item", "")
            display_text = item_text.replace('/', ' / ')
            doc.add_paragraph(f"{i+1}. {display_text}")
            doc.add_paragraph() 

    # Type C: Fill-in, Completion, Ordering, Generic 
    else:
        if qa_pairs:
            if q_type == "fill_in_from_word_bank":
                options = question_data.get("options", [])
                if options:
                    doc.add_paragraph(f"Options: {', '.join(options)}")
            
            if q_type == "match_question_answer":
                doc.add_paragraph("Match the questions and answers:")

            for i, pair in enumerate(qa_pairs):
                item_text = pair.get("item", "")
                if q_type == "match_question_answer":
                    doc.add_paragraph(f"({i+1}) {item_text}")
                    doc.add_paragraph("------------------------------------")
                elif q_type.startswith("order_the_words") or q_type.startswith("create_sentence"):
                    doc.add_paragraph(f"{i+1}. Prompts: {item_text}")
                    doc.add_paragraph("____________________________________________________________________")
                else:
                    doc.add_paragraph(f"{i+1}. {item_text}")
            doc.add_paragraph() # Espaçamento final


def generate_exam_docx(book: str, units: list, questions_config: dict) -> io.BytesIO:
    """
    Generates a .docx document.
    """
    final_doc = Document()
    style = final_doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    try:
        selected_numeric_units = sorted(list(set(re.match(r"(\d+)", u).group(1) for u in units)), key=int)
        units_str = ", ".join(selected_numeric_units)
    except:
        units_str = ", ".join(sorted(units)) 
        
    final_doc.add_heading(get_lang("docx_title").format(book=book, units=units_str), level=0)
    final_doc.add_paragraph(get_lang("docx_name_date"))
    final_doc.add_paragraph()

    final_question_list = []

    numeric_units_map = {}
    for unit_code in units:
        match = re.match(r"(\d+)", unit_code)
        if match:
            num_part = match.group(1)
            if num_part not in numeric_units_map:
                numeric_units_map[num_part] = []
            numeric_units_map[num_part].append(unit_code)

    total_pool_all_questions = []
    book_path = os.path.join(BOOKS_DIR, book)

    if not os.path.exists(book_path):
        return io.BytesIO()

    unit_folders_to_check = [
        f for f in os.listdir(book_path)
        if f.startswith("UNIT-") and f.split('-')[-1] in numeric_units_map.keys()
    ]

    for unit_folder in unit_folders_to_check:
        unit_folder_path = os.path.join(book_path, unit_folder)
        
        for filename in os.listdir(unit_folder_path):
            if filename.endswith(".json"):
                file_path = os.path.join(unit_folder_path, filename)
                
                filename_parts = filename.replace('.json', '').split('-')
                file_section = filename_parts[-1].upper() if len(filename_parts) >= 4 else None

                requested_sections_upper = [s.upper() for s in questions_config.keys() if questions_config[s] > 0]
                if file_section in requested_sections_upper:
                    try:
                        with open(file_path, 'r', encoding='utf-8') as f:
                            data = json.load(f)
                            for q in data.get("questions", []):
                                q['section'] = file_section
                                
                                question_topics = q.get("topic", [])
                                
                                is_relevant = False
                                for req_unit in units: 
                                    if req_unit in question_topics:
                                        is_relevant = True
                                        break
                                
                                if is_relevant:
                                    if 'qa_pair' in q and 'qa_pairs' not in q:
                                        q['qa_pairs'] = q['qa_pair']

                                    total_pool_all_questions.append(q)
                    except Exception:
                        continue

    # 2. Random Selection per Section
    pool_by_section = {s: [] for s in SECTIONS}
    for q in total_pool_all_questions:
        if q['section'] in pool_by_section:
            pool_by_section[q['section']].append(q)

    for section, num_requested in questions_config.items():
        section_upper = section.upper()
        if num_requested > 0 and section_upper in pool_by_section:
            pool = pool_by_section[section_upper]
            num_to_pick = min(num_requested, len(pool))
            if num_to_pick > 0:
                chosen_questions = random.sample(pool, num_to_pick)
                final_question_list.extend(chosen_questions)

    # 3. Write Document and Answer Key
    if not final_question_list:
        final_doc.add_paragraph(get_lang("docx_no_questions_found"))
    else:
        question_counter = 1
        answer_key = []
        current_section = None
        final_question_list.sort(key=lambda q: q['section'])

        for q_data in final_question_list:
            if q_data['section'] != current_section:
                current_section = q_data['section']
                final_doc.add_heading(get_lang("docx_section_header").format(section=current_section.capitalize()), level=1)

            write_question_to_doc(final_doc, q_data, question_counter)

            answers_source = q_data.get("qa_pairs") or q_data.get("qa_pair", [])
            answers_for_this_question = [pair.get("answer", "") for pair in answers_source]
            answer_key.append({"number": question_counter, "answers": answers_for_this_question})
            
            question_counter += 1 

        # Add the Answer Key
        final_doc.add_page_break()
        final_doc.add_heading(get_lang("docx_answer_key_title"), level=1)
        for item in answer_key:
            p_answer_header = final_doc.add_paragraph()
            p_answer_header.add_run(get_lang("docx_answer_key_question").format(number=item['number'])).bold = True
            
            if isinstance(item['answers'], list) and len(item['answers']) > 0:
                for i, answer in enumerate(item['answers']):
                    final_doc.add_paragraph(f"{i+1}. {answer}")
            final_doc.add_paragraph()

    # Save the document to a byte stream in memory
    doc_io = io.BytesIO()
    final_doc.save(doc_io)
    doc_io.seek(0)
    return doc_io

# --- Streamlit Interface ---

# if not os.path.exists(BOOKS_DIR):
#     st.warning(get_lang("warn_books_dir_not_found"))
#     if st.button(get_lang("btn_create_sample_structure")):
#         setup_test_environment()
        
# --- Sidebar for Configuration ---
with st.sidebar:
    st.title(get_lang("sidebar_lang_title"))
    
    # --- SELETOR DE IDIOMA ---
    # Pega o índice do idioma atual para definir o padrão do radio button
    current_lang_code = st.session_state.lang
    lang_codes = list(LANG_OPTIONS_DISPLAY.values())
    default_index = lang_codes.index(current_lang_code) if current_lang_code in lang_codes else 0
    
    selected_lang_display = st.radio(
        label=get_lang("sidebar_lang_label"), 
        options=LANG_OPTIONS_DISPLAY.keys(), 
        index=default_index,
        label_visibility="collapsed" # Esconde o label "Choose language:"
    )
    # Atualiza o session_state se o usuário mudar a seleção
    st.session_state.lang = LANG_OPTIONS_DISPLAY[selected_lang_display]
    
    st.markdown("---") # Divisor
    
    st.header(get_lang("sidebar_header"))

    books = get_available_books(BOOKS_DIR)
    if not books:
        st.error(get_lang("err_no_books").format(BOOKS_DIR))
        st.stop()

    selected_book = st.selectbox(get_lang("sb_select_book"), options=books, index=0)
    parsed_units = parse_available_units(selected_book)

    if not parsed_units:
        st.error(get_lang("err_no_units").format(selected_book))
        st.stop()

    # --- Unit Selection: Refactored to select numeric units only ---
    st.markdown("---")
    st.subheader(get_lang("sb_select_units_title"))

    final_selected_units = []
    
    sorted_numeric_units = sorted(parsed_units.keys(), key=int)
    
    selected_numeric_units = st.multiselect(
        label=get_lang("sb_select_units_label"),
        options=sorted_numeric_units,
        default=sorted_numeric_units,
        key=f"multiselect_{selected_book}_main" 
    )

    for num_unit in selected_numeric_units:
        alpha_units = parsed_units.get(num_unit, [])
        if not alpha_units:
            final_selected_units.append(num_unit)
        else:
            for alpha in alpha_units:
                final_selected_units.append(f"{num_unit}{alpha}")

    # --- End of Unit Selection ---
    st.markdown("---")
    st.subheader(get_lang("sb_q_config_title"))

    questions_config = {
        "grammar": st.number_input(get_lang("sb_q_config_grammar"), min_value=0, max_value=50, value=DEFAULT_QUESTIONS["grammar"], step=1),
        "vocabulary": st.number_input(get_lang("sb_q_config_vocab"), min_value=0, max_value=50, value=DEFAULT_QUESTIONS["vocabulary"], step=1)
    }
    
    total_questions = sum(questions_config.values())

    st.info(get_lang("sb_total_questions").format(total=total_questions))

# --- Main Page Display ---

if not final_selected_units:
    st.warning(get_lang("warn_no_unit_selected"))
    st.stop()

st.title(get_lang("main_summary_title"))
col1, col2, col3 = st.columns(3)

units_display = ", ".join(selected_numeric_units)

with col1:
    st.markdown(get_lang("main_summary_book"))
    st.markdown(f"`{selected_book}`")

with col2:
    st.markdown(get_lang("main_summary_units"))
    st.markdown(f"`{units_display}`")

with col3:
    st.markdown(get_lang("main_summary_q_config"))
    questions_summary = get_lang("main_q_summary_content").format(
        grammar=questions_config['grammar'],
        vocabulary=questions_config['vocabulary']
    )
    st.markdown(f"`{questions_summary}`")

# Determine button text based on whether the configuration is default or custom
is_default_config = (
    questions_config["grammar"] == DEFAULT_QUESTIONS["grammar"] and
    questions_config["vocabulary"] == DEFAULT_QUESTIONS["vocabulary"]
)
button_text = get_lang("btn_generate_std") if is_default_config else get_lang("btn_generate_custom")

# Initialize session state for storing generated exam data
if 'exam_data' not in st.session_state:
    st.session_state.exam_data = None
if 'exam_filename' not in st.session_state:
    st.session_state.exam_filename = 'test.docx'

if st.button(button_text, type="primary", use_container_width=True, disabled=(total_questions == 0)):
    if total_questions > 0:
        with st.spinner(get_lang("spinner_generating")):
            try:
                exam_bytes = generate_exam_docx(selected_book, final_selected_units, questions_config) 
                
                st.session_state.exam_data = exam_bytes
                
                numeric_units_filename = "_".join(selected_numeric_units)
                st.session_state.exam_filename = get_lang("filename_test").format(
                    book=selected_book, 
                    units=numeric_units_filename
                )
            except Exception as e:
                st.error(get_lang("err_generation").format(error=e))
                st.session_state.exam_data = None
    else:
        st.warning(get_lang("warn_no_questions_selected"))

if st.session_state.exam_data:
    st.markdown("---")
    st.download_button(
        label=get_lang("btn_download"),
        data=st.session_state.exam_data,
        file_name=st.session_state.get('exam_filename', 'test.docx'),
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        use_container_width=True
    )

st.markdown("---")
st.title(get_lang("about_title"))

st.write(get_lang("about_p1"))
st.write(get_lang("about_p2"))

st.markdown("---")

st.title(get_lang("contribute_title"))

st.markdown(get_lang("contribute_link_text").format(link=EVALUATION_LINK))

st.markdown(f"""
1. {get_lang("contribute_li1")}
2. {get_lang("contribute_li2")}
3. {get_lang("contribute_li3")}
""")

st.markdown("---")
st.markdown(get_lang("footer_text"))